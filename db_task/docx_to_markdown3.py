import os
import re
import logging
import docx
from typing import Dict, Any, List, Tuple

logger = logging.getLogger(__name__)

class DocxToMarkdown:
    """
    将DOCX文档转换为Markdown格式
    """
    
    def __init__(self):
        """初始化转换器"""
        pass
    
    def convert(self, file_path: str) -> str:
        """
        将DOCX文件转换为Markdown格式
        
        参数:
            file_path (str): DOCX文件路径
            
        返回:
            str: 转换后的Markdown文本
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            if not file_path.lower().endswith('.docx'):
                raise ValueError(f"不支持的文件格式，仅支持.docx文件: {file_path}")
            
            document = docx.Document(file_path)
            markdown_text = self._process_document(document)
            
            return markdown_text
            
        except Exception as e:
            logger.error(f"转换DOCX到Markdown失败: {str(e)}")
            raise ValueError(f"转换失败: {str(e)}")
    
    def _process_document(self, document) -> str:
        """
        处理文档对象并转换为Markdown
        
        参数:
            document: docx文档对象
            
        返回:
            str: 转换后的Markdown文本
        """
        markdown_parts = []
        
        # 处理段落
        for para in document.paragraphs:
            if not para.text.strip():
                continue
            
            # 跳过目录行（包含数字+标题+页码的模式）
            if re.match(r'^(目录|[0-9]+(\.[0-9]+)*)\s+.*\s+\d+$', para.text.strip()):
                continue
                
            # 获取段落的大纲级别，修复 CT_PPr 对象没有 get_or_add_outlineLvl 属性的问题
            try:
                # 尝试获取段落样式
                if para.style and para.style.name.startswith('Heading'):
                    # 从标题样式名称中提取级别（如 Heading 1 -> 1）
                    level_match = re.search(r'Heading\s+(\d+)', para.style.name)
                    level_value = int(level_match.group(1)) if level_match else None
                    
                    if level_value is not None:
                        # 这是标题段落
                        markdown_parts.append(f"{'#' * level_value} {para.text}")
                        continue
            except Exception as e:
                logger.warning(f"处理段落样式时出错: {str(e)}")
            
            # 检查是否为目录项（如"密码应用需求分析"这样的无标记目录项）
            if para.text.strip() and not para.text.startswith('#'):
                # 跳过目录标题
                if para.text.strip() == "目录":
                    continue
                    
                # 检查是否为目录项（包含数字编号和页码）
                if re.match(r'^[0-9]+(\.[0-9]+)*\s+.*\s+\d+$', para.text.strip()):
                    continue
                    
                # 检查是否为可能的目录项（简短的单行文本，没有标点符号）
                if len(para.text) < 50 and '\n' not in para.text and not re.search(r'[,.;:!?，。；：！？]', para.text):
                    # 将其视为二级标题
                    markdown_parts.append(f"## {para.text}")
                    continue
            
            # 处理普通段落
            text = para.text
            
            # 处理加粗和斜体
            for run in para.runs:
                if run.bold and run.italic:
                    text = text.replace(run.text, f"***{run.text}***")
                elif run.bold:
                    text = text.replace(run.text, f"**{run.text}**")
                elif run.italic:
                    text = text.replace(run.text, f"*{run.text}*")
            
            markdown_parts.append(text)
        
        # 处理表格
        for table in document.tables:
            table_rows = []
            
            # 表头
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            table_rows.append("| " + " | ".join(header_row) + " |")
            
            # 分隔行
            separator = "| " + " | ".join(["---"] * len(header_row)) + " |"
            table_rows.append(separator)
            
            # 数据行
            for row in table.rows[1:]:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text.strip())
                table_rows.append("| " + " | ".join(row_cells) + " |")
            
            markdown_parts.append("\n".join(table_rows))
        
        # 合并所有内容
        return "\n\n".join(markdown_parts)
    
    def convert_and_chunk(self, file_path: str) -> List[Dict[str, Any]]:
        """
        将DOCX文件转换为Markdown格式并按标题分块
        
        参数:
            file_path (str): DOCX文件路径
            
        返回:
            List[Dict[str, Any]]: 按标题分块的内容列表，每块包含内容和所属章节信息
        """
        markdown_text = self.convert(file_path)
        return self._chunk_by_headers(markdown_text)
    
    def _chunk_by_headers(self, markdown_text: str) -> List[Dict[str, Any]]:
        """
        按标题将Markdown文本分块
        
        参数:
            markdown_text (str): Markdown文本
            
        返回:
            List[Dict[str, Any]]: 分块后的内容列表
        """
        lines = markdown_text.split('\n')
        chunks = []
        
        current_h1 = ""  # 当前一级标题
        current_h2 = ""  # 当前二级标题
        current_chunk = {"content": [], "h1": "", "h2": "", "title": ""}
        
        for line in lines:
            if not line.strip():
                continue
                
            # 检查是否为标题行
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                # 如果当前块有内容，保存它
                if current_chunk["content"]:
                    chunks.append({
                        "content": "\n".join(current_chunk["content"]),
                        "h1": current_chunk["h1"],
                        "h2": current_chunk["h2"],
                        "title": current_chunk["title"]
                    })
                
                level = len(header_match.group(1))
                title_text = header_match.group(2)
                
                # 更新当前标题级别
                if level == 1:
                    current_h1 = title_text
                    current_h2 = ""
                elif level == 2:
                    current_h2 = title_text
                
                # 创建新块
                current_chunk = {
                    "content": [line],
                    "h1": current_h1,
                    "h2": current_h2,
                    "title": title_text
                }
            else:
                # 将行添加到当前块
                current_chunk["content"].append(line)
        
        # 添加最后一个块
        if current_chunk["content"]:
            chunks.append({
                "content": "\n".join(current_chunk["content"]),
                "h1": current_chunk["h1"],
                "h2": current_chunk["h2"],
                "title": current_chunk["title"]
            })
        
        return chunks


if __name__ == "__main__":
    docx_to_markdown = DocxToMarkdown()
    # markdown_text = docx_to_markdown.convert(r""D:\code\llm_sass_station\EasyRAG\gov\贵州省政务云密码服务与监管平台建设方案v2.5 by wbl.docx"")
    # print(markdown_text)
    
    # 按标题分块并获取章节信息
    chunks = docx_to_markdown.convert_and_chunk(r"D:\code\llm_sass_station\EasyRAG\gov\贵州省政务云密码服务与监管平台建设方案v2.docx")
    print(f"\n总共分成了 {len(chunks)} 个块")
    for i, chunk in enumerate(chunks):  # 只打印前3个块作为示例
        print('********************************************')
        print(f"\n块 {i+1}:")
        print(f"一级标题: {chunk['h1']}")
        print(f"二级标题: {chunk['h2']}")
        print(f"当前标题: {chunk['title']}")
        print(f"内容预览: {chunk['content']}...")
