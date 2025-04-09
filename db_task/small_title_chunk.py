import os
import traceback
import re
import logging
import uuid
from typing import Dict, List, Tuple, Union, Optional, Any
from urllib.parse import quote
import docx
from docx import Document
import fitz  # PyMuPDF
from fastapi import HTTPException
import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
# from core.file_read.ocr_extract import extract_text_with_subprocess

logger = logging.getLogger(__name__)

"""create by haozi
    2025-03-14
    文件处理类，支持从不同格式的文件中提取内容
"""


class FileHandler:
    """文件处理类，支持从不同格式的文件中提取内容"""
    
    def __init__(self):
        """初始化文件处理器"""
        pass
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        处理文件并返回提取的文本内容及文件结构
        
        参数:
            file_path (str): 文件路径
            
        返回:
            Dict[str, Any]: 包含文本内容和结构信息的字典
        """
        try:
            # 检查文件是否存在和可访问
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                raise ValueError(f"文件为空: {file_path}")
                
            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)
            
            print(f"正在处理文件: {file_path}, 类型: {file_ext}, 大小: {file_size} 字节")
            
            result = {
                "file_name": file_name,
                "file_path": file_path,
                "file_type": file_ext.replace('.', ''),
                "content": "",
                "structure": {},
                "paragraphs": [],  # 添加段落结构
                "sections": []  # 添加按小标题分割的章节
            }
            
            try:
                if file_ext == '.docx':
                    content, structure, sections = self._read_docx_file_with_sections(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                    result["sections"] = sections
                else:
                    message = f"不支持的文件格式: {file_ext}"
                    # print(message)
                    result["error"] = message
                    result["content"] = f"无法处理此文件格式: {file_ext}"
            except Exception as e:
                # 捕获并记录特定文件处理错误
                error_message = f"处理文件内容时出错: {str(e)}"
                print(error_message)
                traceback.print_exc()
                result["error"] = error_message
                result["content"] = f"文件处理错误: {str(e)}"
            
            # 确保始终返回有效内容，即使是错误消息
            if not result["content"]:
                result["content"] = "未能提取文件内容"
                
            return result
            
        except Exception as e:
            # 捕获顶层异常
            error_message = f"处理文件时出错: {str(e)}"
            print(error_message)
            traceback.print_exc()
            return {
                "file_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
                "content": f"文件处理失败: {str(e)}",
                "structure": {},
                "paragraphs": [],
                "error": str(e)
            }
    
    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本内容分割成段落结构，确保表格和图片与其上下文保持在一起
        
        参数:
            text (str): 要分割的文本内容
            
        返回:
            List[Dict[str, Any]]: 段落列表，每个段落包含内容和位置信息
        """
        paragraphs = []
        
        # 第一步：检测所有表格和图片位置
        table_pattern = re.compile(r'\n表格\s+\d+:[\s\S]+?(?=\n\n表格\s+\d+:|$)')
        image_pattern = re.compile(r'\n插图开始[\s\S]+?插图结束\n')
        
        table_matches = list(table_pattern.finditer(text))
        image_matches = list(image_pattern.finditer(text))
        
        # 收集所有特殊块的位置（表格和图片）
        special_blocks = []
        for m in table_matches:
            special_blocks.append(("table", m.start(), m.end()))
        for m in image_matches:
            special_blocks.append(("image", m.start(), m.end()))
        
        # 按照在文档中的位置排序
        special_blocks.sort(key=lambda x: x[1])
        
        # 第二步：查找这些特殊块的上下文段落
        context_blocks = []
        
        # 确定段落边界
        all_paragraphs = re.split(r'\n\s*\n', text)
        paragraph_positions = []
        
        pos = 0
        for para in all_paragraphs:
            start = pos
            end = pos + len(para)
            if para.strip():  # 只记录非空段落
                paragraph_positions.append((start, end))
            pos = end + 2  # +2 是为了考虑分隔符 \n\n
        
        # 为每个特殊块找到其上下文段落
        for block_type, block_start, block_end in special_blocks:
            # 查找块前的相关段落作为上下文
            context_start = 0
            for para_start, para_end in paragraph_positions:
                if para_end < block_start and para_end > context_start:
                    # 检查段落是否与特殊块相关
                    para_text = text[para_start:para_end]
                    
                    # 表格相关：包含"表"字或距离较近
                    if block_type == "table" and ("表" in para_text or (block_start - para_end) < 100):
                        context_start = para_start
                    
                    # 图片相关：包含"图"字或"如下所示"等描述性词语或距离较近
                    elif block_type == "image" and ("图" in para_text or 
                                                  "如下所示" in para_text or 
                                                  "示意图" in para_text or
                                                  (block_start - para_end) < 100):
                        context_start = para_start
            
            if context_start > 0:
                context_blocks.append((context_start, block_end, block_type))
            else:
                context_blocks.append((block_start, block_end, block_type))
        
        # 合并重叠的上下文块（可能一段文字同时描述了表格和图片）
        merged_blocks = []
        if context_blocks:
            # 按起始位置排序
            context_blocks.sort(key=lambda x: x[0])
            current_block = context_blocks[0]
            
            for next_start, next_end, next_type in context_blocks[1:]:
                curr_start, curr_end, curr_type = current_block
                
                # 如果有重叠，合并块
                if next_start <= curr_end:
                    current_block = (curr_start, max(curr_end, next_end), f"{curr_type}_{next_type}")
                else:
                    merged_blocks.append(current_block)
                    current_block = (next_start, next_end, next_type)
            
            merged_blocks.append(current_block)
        
        # 第三步：处理常规文本和特殊块
        processed_ranges = []
        for start, end, block_type in merged_blocks:
            processed_ranges.append((start, end))
        
        # 处理未被特殊块覆盖的文本
        pos = 0
        while pos < len(text):
            # 检查当前位置是否已处理
            inside_processed = False
            for start, end in processed_ranges:
                if start <= pos < end:
                    inside_processed = True
                    pos = end
                    break
            
            if inside_processed:
                continue
            
            # 找到下一个已处理区域
            next_start = len(text)
            for start, end in processed_ranges:
                if start > pos and start < next_start:
                    next_start = start
            
            # 处理这段普通文本
            if next_start > pos:
                segment_text = text[pos:next_start].strip()
                if segment_text:
                    regular_paragraphs = self._process_regular_text(segment_text, pos)
                    paragraphs.extend(regular_paragraphs)
                pos = next_start
        
        # 添加特殊块（表格/图片及其上下文）
        for idx, (start, end, block_type) in enumerate(merged_blocks):
            block_text = text[start:end].strip()
            if block_text:
                # 根据块类型设置不同的标记
                contains_table = "table" in block_type
                contains_image = "image" in block_type
                
                paragraphs.append({
                    "id": f"special_block_{idx}",
                    "content": block_text,
                    "position": start,
                    "length": end - start,
                    "is_heading": False,
                    "heading_level": 0,
                    "contains_table": contains_table,
                    "contains_image": contains_image
                })
        
        # 按位置排序段落
        paragraphs.sort(key=lambda x: x["position"])
        
        return paragraphs
    
    def _process_regular_text(self, text: str, base_position: int = 0) -> List[Dict[str, Any]]:
        """
        处理普通文本（非表格相关）并分割成段落
        
        参数:
            text (str): 要处理的文本
            base_position (int): 文本在原始文档中的起始位置
            
        返回:
            List[Dict[str, Any]]: 处理后的段落列表
        """
        result = []
        raw_paragraphs = text.split('\n\n')
        
        start_pos = base_position
        for i, para in enumerate(raw_paragraphs):
            if para.strip():  # 忽略空段落
                # 检测是否为标题
                # print(i,  para)
                is_heading = False
                heading_level = 0
                header_match = re.match(r'^(#{1,6})\s+(.+)$', para.strip())
                if header_match:
                    is_heading = True
                    heading_level = len(header_match.group(1))
                
                paragraph = {
                    "id": f"p_{start_pos}",
                    "content": para.strip(),
                    "position": start_pos,
                    "length": len(para),
                    "is_heading": is_heading,
                    "heading_level": heading_level if is_heading else 0,
                    "contains_table": False,
                    "contains_image": False  # 添加图片标记
                }
                result.append(paragraph)
            
            start_pos += len(para) + 2  # +2 是为了考虑分隔符 \n\n
        
        return result
    
    def _read_docx_file_with_sections(self, file_path: str) -> Tuple[str, Dict, List[Dict]]:
        """读取docx文件，保留标题级别信息，并识别小标题"""
        document = None
        try:
            document = docx.Document(file_path)
            
            # 将Word转为Markdown保留标题
            markdown_text = ""
            for para in document.paragraphs:
                if para.style.name == 'Heading 1':
                    markdown_text += f"# {para.text}\n\n"
                elif para.style.name == 'Heading 2':
                    markdown_text += f"## {para.text}\n\n"
                else:
                    markdown_text += f"{para.text}\n\n"
            
            # 然后使用SubheadingTextSplitter处理
            content, structure, sections = self._read_docx_file_with_sections(markdown_text)
            
            return content, structure, sections
        except Exception as e:
            logger.error(f"读取docx文件失败: {str(e)}")
            raise ValueError(f"无法处理docx文件: {str(e)}")

    def _read_docx_file(self, file_path: str) -> Tuple[str, Dict]:
        """读取docx文件，保留标题级别信息"""
        document = None
        try:
            document = docx.Document(file_path)
            
            # 提取段落文本，为标题添加#标记
            formatted_paragraphs = []
            for para in document.paragraphs:
                if para.text.strip():
                    # 检查段落是否为标题
                    if para.style.name.startswith('Heading'):
                        # 提取标题级别（Heading 1, Heading 2等）
                        try:
                            heading_level = int(para.style.name.split()[-1])
                            # 添加对应数量的#作为标题标记
                            formatted_text = '#' * heading_level + ' ' + para.text
                            formatted_paragraphs.append(formatted_text)
                        except ValueError:
                            # 如果无法解析标题级别，按普通段落处理
                            formatted_paragraphs.append(para.text)
                    else:
                        formatted_paragraphs.append(para.text)
            
            # 合并所有段落为一个字符串
            content = "\n\n".join(formatted_paragraphs)
            
            # 创建结构化信息
            structure = {
                "paragraphs": formatted_paragraphs,
                "tables": []
            }
            
            # 提取表格内容
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                structure["tables"].append(table_data)
                
            return content, structure
        except Exception as e:
            logger.error(f"读取docx文件失败: {str(e)}")
            raise ValueError(f"无法处理docx文件: {str(e)}")

    def read_file(self, file_path: str) -> Dict[str, Any]:
        """
        读取文件并提取内容
        
        参数:
            file_path (str): 文件路径
            
        返回:
            Dict[str, Any]: 包含文本内容和结构信息的字典
        """
        result = {
            "content": "",
            "structure": {},
            "paragraphs": []
        }
        
        try:
            # 检查文件是否存在
            print(f"开始读取文件: {file_path}")
            if not os.path.exists(file_path):
                error_msg = f"文件不存在: {file_path}"
                print(error_msg)
                raise FileNotFoundError(error_msg)
                
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            print(f"文件大小: {file_size} 字节")
            if file_size == 0:
                error_msg = f"文件为空: {file_path}"
                print(error_msg)
                raise ValueError(error_msg)
                
            # 获取文件扩展名
            _, file_ext = os.path.splitext(file_path.lower())
            print(f"文件扩展名: {file_ext}")
            
            try:
                if file_ext == '.docx':
                    print(f"处理Word文档(docx)...")
                    content, structure, sections = self._read_docx_file_with_sections(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                    result["sections"] = sections
                else:
                    # 检查文件类型
                    import mimetypes
                    mime_type, _ = mimetypes.guess_type(file_path)
                    print(f"未知文件类型: {file_ext}, MIME类型: {mime_type}")
                    
                    error_msg = f"不支持的文件类型: {file_ext}"
                    print(error_msg)
                    raise ValueError(error_msg)
            except Exception as e:
                print(f"处理文件时出错: {str(e)}")
                traceback.print_exc()
                raise
                
            print(f"文件处理成功，提取内容长度: {len(result['content'])}")
            if not result["content"]:
                print(f"警告: 提取的内容为空")
                
            return result
        except Exception as e:
            print(f"读取文件失败: {str(e)}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"读取文件失败: {str(e)}")

if __name__ == "__main__":
    file_path = r"C:\Users\Administrator\Desktop\培训资料\Docs\技术方案\密码应用方案\密码应用集成方案.docx"
    handler = FileHandler()
    result = handler.process_file(file_path)
    # print(result.keys())
    # print(result)
    # for line in result["structure"]["paragraphs"]:
    #     print(1, line)
    # print(result["structure"]["sections"])


# if __name__ == "__main__":
#     # 测试代码
#     chunker = FileHandler()
#     file_path = r"C:\Users\Administrator\Desktop\培训资料\Docs\技术方案\密码应用方案\密码应用集成方案.docx"  # 替换为实际的docx文件路径
#     chunks = chunker.process_docx(file_path)
#     print(chunks)
