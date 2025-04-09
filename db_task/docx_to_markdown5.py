import os
import re
import logging
import docx
from typing import Dict, Any, List, Tuple
import cv2
import tempfile
import sys
import json
import subprocess
import traceback
import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.file_read.ocr_extract import extract_text_with_subprocess

logger = logging.getLogger(__name__)

class DocxToMarkdown:
    """
    将DOCX文档转换为Markdown格式
    """
    
    def __init__(self):
        """初始化转换器"""
        pass
    
    def convert(self, file_path: str) -> str:
        """
        将DOCX文件转换为Markdown格式
        
        参数:
            file_path (str): DOCX文件路径
            
        返回:
            str: 转换后的Markdown文本
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            if not file_path.lower().endswith('.docx'):
                raise ValueError(f"不支持的文件格式，仅支持.docx文件: {file_path}")
            
            document = docx.Document(file_path)
            markdown_text = self._process_document(document)
            # print(markdown_text)
            
            return markdown_text
            
        except Exception as e:
            logger.error(f"转换DOCX到Markdown失败: {str(e)}")
            raise ValueError(f"转换失败: {str(e)}")
    
    def _process_document(self, document) -> str:
        """
        处理文档对象并转换为Markdown
        
        参数:
            document: docx文档对象
            
        返回:
            str: 转换后的Markdown文本
        """
        # 构建内容元素列表，包括段落和表格的混合顺序
        document_elements = []
        
        # 获取所有元素（段落和表格）
        # 注意：python-docx没有直接提供获取元素在文档中顺序的方法
        # 我们需要从document._element中提取
        try:
            # 尝试通过document._element获取所有内容元素
            all_elements = []
            for element in document._element.body.iterchildren():
                if element.tag.endswith(('p', 'tbl')):
                    all_elements.append(element)
            
            # 将元素映射到对应的段落或表格对象
            for element in all_elements:
                if element.tag.endswith('p'):
                    # 找到对应的段落对象
                    for para in document.paragraphs:
                        if para._element is element:
                            document_elements.append(('paragraph', para))
                            break
                elif element.tag.endswith('tbl'):
                    # 找到对应的表格对象
                    for table in document.tables:
                        if table._element is element:
                            document_elements.append(('table', table))
                            break
        except Exception as e:
            logger.warning(f"无法按顺序提取文档元素，将分开处理段落和表格: {str(e)}")
            # 回退到分别处理段落和表格
            for para in document.paragraphs:
                document_elements.append(('paragraph', para))
            for table in document.tables:
                document_elements.append(('table', table))
        
        # 如果没有成功提取元素，则回退到原来的方式
        if not document_elements:
            logger.warning("未能提取到任何文档元素，将分开处理段落和表格")
            for para in document.paragraphs:
                document_elements.append(('paragraph', para))
            for table in document.tables:
                document_elements.append(('table', table))
        
        # 处理所有元素
        markdown_parts = []
        
        for element_type, element in document_elements:
            if element_type == 'paragraph':
                para = element
                if not para.text.strip():
                    continue
                
                # 跳过目录行（包含数字+标题+页码的模式）
                if re.match(r'^(目录|[0-9]+(\.[0-9]+)*)\s+.*\s+\d+$', para.text.strip()):
                    continue
                    
                # 获取段落的大纲级别，修复 CT_PPr 对象没有 get_or_add_outlineLvl 属性的问题
                try:
                    # 尝试获取段落样式
                    if para.style and para.style.name.startswith('Heading'):
                        # 从标题样式名称中提取级别（如 Heading 1 -> 1）
                        level_match = re.search(r'Heading\s+(\d+)', para.style.name)
                        level_value = int(level_match.group(1)) if level_match else None
                        
                        if level_value is not None:
                            # 这是标题段落
                            markdown_parts.append(f"{'#' * level_value} {para.text}")
                            continue
                except Exception as e:
                    logger.warning(f"处理段落样式时出错: {str(e)}")
                
                # 检查是否为目录项（如"密码应用需求分析"这样的无标记目录项）
                if para.text.strip() and not para.text.startswith('#'):
                    # 跳过目录标题
                    if para.text.strip() == "目录":
                        continue
                        
                    # 检查是否为目录项（包含数字编号和页码）
                    if re.match(r'^[0-9]+(\.[0-9]+)*\s+.*\s+\d+$', para.text.strip()):
                        continue
                        
                    # 检查是否为可能的目录项（简短的单行文本，没有标点符号）
                    if len(para.text) < 50 and '\n' not in para.text and not re.search(r'[,.;:!?，。；：！？]', para.text):
                        # 将其视为二级标题
                        markdown_parts.append(f"## {para.text}")
                        continue
                
                # 检查是否为伪表格（使用空格或制表符对齐的文本）
                if self._is_pseudo_table(para.text):
                    # 将伪表格转换为Markdown表格
                    table_md = self._convert_pseudo_table(para.text)
                    markdown_parts.append(table_md)
                    continue
                
                # 处理普通段落
                text = para.text
                
                # 处理加粗和斜体
                for run in para.runs:
                    if run.bold and run.italic:
                        text = text.replace(run.text, f"***{run.text}***")
                    elif run.bold:
                        text = text.replace(run.text, f"**{run.text}**")
                    elif run.italic:
                        text = text.replace(run.text, f"*{run.text}*")
                
                markdown_parts.append(text)
                
            elif element_type == 'table':
                table = element
                # 处理表格
                # 尝试使用OCR提取表格内容
                try:
                    table_image = self._extract_table_image(table)
                    if table_image:
                        # 使用OCR提取表格内容
                        table_text = extract_text_with_subprocess(table_image)
                        if table_text:
                            markdown_parts.append(f"```\n{table_text}\n```")
                            continue
                except Exception as e:
                    logger.error(f"使用OCR提取表格失败: {str(e)}")
                    logger.error(traceback.format_exc())
                
                # 如果OCR提取失败，回退到常规表格处理
                table_rows = []
                
                # 确保表格至少有一行
                if len(table.rows) > 0:
                    # 表头
                    header_row = []
                    for cell in table.rows[0].cells:
                        header_row.append(cell.text.strip())
                    table_rows.append("| " + " | ".join(header_row) + " |")
                    
                    # 分隔行
                    separator = "| " + " | ".join(["---"] * len(header_row)) + " |"
                    table_rows.append(separator)
                    
                    # 数据行
                    for row in table.rows[1:]:
                        row_cells = []
                        for cell in row.cells:
                            row_cells.append(cell.text.strip())
                        table_rows.append("| " + " | ".join(row_cells) + " |")
                
                    markdown_parts.append("\n".join(table_rows))
        
        # 合并所有内容
        return "\n\n".join(markdown_parts)
    
    def _is_pseudo_table(self, text: str) -> bool:
        """
        检测文本是否为伪表格（使用空格或制表符对齐的文本）
        
        参数:
            text: 要检查的文本
            
        返回:
            bool: 是否为伪表格
        """
        lines = text.split('\n')
        if len(lines) < 2:
            return False
            
        # 检查是否有多行且每行都有类似的空格模式
        space_patterns = []
        for line in lines:
            # 查找连续空格或制表符的位置
            spaces = [match.start() for match in re.finditer(r'\s{2,}|\t+', line)]
            if spaces:
                space_patterns.append(spaces)
                
        # 如果至少有两行且空格模式相似，则可能是伪表格
        if len(space_patterns) >= 2:
            # 检查空格模式的相似性
            pattern_lengths = [len(pattern) for pattern in space_patterns]
            if max(pattern_lengths) > 0 and min(pattern_lengths) > 0:
                return True
                
        return False
    
    def _convert_pseudo_table(self, text: str) -> str:
        """
        将伪表格文本转换为Markdown表格
        
        参数:
            text: 伪表格文本
            
        返回:
            str: Markdown表格
        """
        lines = text.split('\n')
        if not lines:
            return text
            
        # 使用正则表达式分割每行
        table_data = []
        for line in lines:
            # 使用连续空格或制表符分割
            cells = re.split(r'\s{2,}|\t+', line.strip())
            cells = [cell.strip() for cell in cells if cell.strip()]
            if cells:
                table_data.append(cells)
                
        if not table_data:
            return text
            
        # 确定表格的最大列数
        max_cols = max(len(row) for row in table_data)
        
        # 创建Markdown表格
        md_table = []
        
        # 添加表头
        header = table_data[0]
        # 确保表头有足够的列
        while len(header) < max_cols:
            header.append("")
        md_table.append("| " + " | ".join(header) + " |")
        
        # 添加分隔行
        md_table.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        # 添加数据行
        for row in table_data[1:]:
            # 确保每行有足够的列
            while len(row) < max_cols:
                row.append("")
            md_table.append("| " + " | ".join(row) + " |")
            
        return "\n".join(md_table)
    
    def _extract_table_image(self, table):
        """
        从表格对象提取图像
        
        参数:
            table: docx表格对象
            
        返回:
            str: 临时图像文件路径或None
        """
        try:
            # 这里需要实现表格转图像的逻辑
            # 由于python-docx不直接支持表格转图像，这里是一个占位
            # 实际实现可能需要使用其他库或方法
            return None
        except Exception as e:
            logger.error(f"提取表格图像失败: {str(e)}")
            return None
    
    def convert_and_chunk(self, file_path: str) -> List[Dict[str, Any]]:
        """
        将DOCX文件转换为Markdown格式并按标题分块
        
        参数:
            file_path (str): DOCX文件路径
            
        返回:
            List[Dict[str, Any]]: 按标题分块的内容列表，每块包含内容和所属章节信息
        """
        markdown_text = self.convert(file_path)
        return self._chunk_by_headers(markdown_text)
    
    def _chunk_by_headers(self, markdown_text: str) -> List[Dict[str, Any]]:
        """
        按标题将Markdown文本分块，确保表格与周围段落内容保持在同一分块中
        
        参数:
            markdown_text (str): Markdown文本
            
        返回:
            List[Dict[str, Any]]: 分块后的内容列表
        """
        lines = markdown_text.split('\n')
        chunks = []
        
        current_h1 = ""  # 当前一级标题
        current_h2 = ""  # 当前二级标题
        current_chunk = {"content": [], "h1": "", "h2": "", "title": ""}
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
                
            # 检查是否为标题行
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                # 如果当前块有内容，保存它
                if current_chunk["content"]:
                    chunks.append({
                        "content": "\n".join(current_chunk["content"]),
                        "h1": current_chunk["h1"],
                        "h2": current_chunk["h2"],
                        "title": current_chunk["title"]
                    })
                
                level = len(header_match.group(1))
                title_text = header_match.group(2)
                
                # 更新当前标题级别
                if level == 1:
                    current_h1 = title_text
                    current_h2 = ""
                elif level == 2:
                    current_h2 = title_text
                
                # 创建新块
                current_chunk = {
                    "content": [lines[i]],
                    "h1": current_h1,
                    "h2": current_h2,
                    "title": title_text
                }
                i += 1
            else:
                # 处理表格
                if line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
                    # 将整个表格作为一个整体添加到当前块
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        table_lines.append(lines[i])
                        i += 1
                    current_chunk["content"].extend(table_lines)
                # 处理代码块中的表格 (```包裹的内容)
                elif line.startswith('```'):
                    code_block_lines = [lines[i]]
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_block_lines.append(lines[i])
                        i += 1
                    if i < len(lines):  # 添加结束的 ```
                        code_block_lines.append(lines[i])
                        i += 1
                    current_chunk["content"].extend(code_block_lines)
                else:
                    # 普通行，添加到当前块
                    current_chunk["content"].append(lines[i])
                    i += 1
        
        # 添加最后一个块
        if current_chunk["content"]:
            chunks.append({
                "content": "\n".join(current_chunk["content"]),
                "h1": current_chunk["h1"],
                "h2": current_chunk["h2"],
                "title": current_chunk["title"]
            })
        
        return chunks


if __name__ == "__main__":
    docx_to_markdown = DocxToMarkdown()
    # markdown_text = docx_to_markdown.convert(r""D:\code\llm_sass_station\EasyRAG\gov\贵州省政务云密码服务与监管平台建设方案v2.5 by wbl.docx"")
    # print(markdown_text)
    
    # 按标题分块并获取章节信息
    chunks = docx_to_markdown.convert_and_chunk(r"D:\code\llm_sass_station\EasyRAG\gov\贵州省政务云密码服务与监管平台建设方案v2.docx")
    print(f"\n总共分成了 {len(chunks)} 个块")
    for i, chunk in enumerate(chunks):  # 只打印前3个块作为示例
        print('********************************************')
        print(f"\n块 {i+1}:")
        print(f"一级标题: {chunk['h1']}")
        print(f"二级标题: {chunk['h2']}")
        print(f"当前标题: {chunk['title']}")
        print(f"内容预览: {chunk['content']}...")
