import os
import traceback
import re
import logging
import uuid
from typing import Dict, List, Tuple, Union, Optional, Any
from urllib.parse import quote
import docx
from docx import Document
import fitz  # PyMuPDF
from fastapi import HTTPException
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from core.file_read.ocr_extract import extract_text_with_subprocess

logger = logging.getLogger(__name__)

"""create by haozi
    2025-03-14
    文件处理类，支持从不同格式的文件中提取内容并转换为Markdown格式
"""


class FileToMarkdown:
    """文件处理类，支持从不同格式的文件中提取内容并转换为Markdown格式"""
    
    def __init__(self):
        """初始化文件处理器"""
        pass
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        处理文件并返回提取的文本内容及文件结构
        
        参数:
            file_path (str): 文件路径
            
        返回:
            Dict[str, Any]: 包含文本内容和结构信息的字典
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)
            
            logger.info(f"开始处理文件: {file_path}, 类型: {file_ext}")
            
            result = {
                "file_name": file_name,
                "file_path": file_path,
                "file_type": file_ext.replace('.', ''),
                "content": "",
                "structure": {},
                "paragraphs": []  # 添加段落结构
            }
            
            try:
                # 根据文件类型调用不同的转换方法
                if file_ext == '.txt':
                    result["content"] = self._txt_to_markdown(file_path)
                    result["paragraphs"] = self._split_into_paragraphs(result["content"])
                elif file_ext == '.docx':
                    markdown_content = self._docx_to_markdown(file_path)
                    result["content"] = markdown_content
                    result["structure"] = {"format": "markdown"}
                    result["paragraphs"] = self._split_into_paragraphs(markdown_content)
                elif file_ext == '.doc':
                    markdown_content = self._doc_to_markdown(file_path)
                    result["content"] = markdown_content
                    result["structure"] = {"format": "markdown"}
                    result["paragraphs"] = self._split_into_paragraphs(markdown_content)
                elif file_ext == '.pdf':
                    markdown_content = self._pdf_to_markdown(file_path)
                    result["content"] = markdown_content
                    result["structure"] = {"format": "markdown"}
                    result["paragraphs"] = self._split_into_paragraphs(markdown_content)
                else:
                    message = f"不支持的文件格式: {file_ext}"
                    logger.error(message)
                    result["error"] = message
                    result["content"] = f"无法处理此文件格式: {file_ext}"
            except Exception as e:
                # 捕获并记录特定文件处理错误
                error_message = f"处理文件内容时出错: {str(e)}"
                logger.error(error_message)
                logger.error(traceback.format_exc())
                result["error"] = error_message
                result["content"] = f"文件处理错误: {str(e)}"
            
            # 确保始终返回有效内容，即使是错误消息
            if not result["content"]:
                result["content"] = "未能提取文件内容"
                
            return result
            
        except Exception as e:
            # 捕获顶层异常
            error_message = f"处理文件时出错: {str(e)}"
            logger.error(error_message)
            logger.error(traceback.format_exc())
            return {
                "file_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
                "content": f"文件处理失败: {str(e)}",
                "structure": {},
                "paragraphs": [],
                "error": str(e)
            }
    
    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本内容分割成段落结构
        
        参数:
            text (str): 要分割的文本内容
            
        返回:
            List[Dict[str, Any]]: 段落列表，每个段落包含内容和位置信息
        """
        paragraphs = []
        
        # 分割段落
        raw_paragraphs = re.split(r'\n\s*\n', text)
        
        start_pos = 0
        for i, para in enumerate(raw_paragraphs):
            if para.strip():  # 忽略空段落
                # 检测是否为标题
                is_heading = False
                heading_level = 0
                header_match = re.match(r'^(#{1,6})\s+(.+)$', para.strip())
                if header_match:
                    is_heading = True
                    heading_level = len(header_match.group(1))
                
                # 检测是否包含表格
                contains_table = bool(re.search(r'\|[-\s|]+\|', para))
                
                # 检测是否包含图片
                contains_image = bool(re.search(r'!\[.*?\]\(.*?\)', para) or '插图开始' in para)
                
                paragraph = {
                    "id": f"p_{i}",
                    "content": para.strip(),
                    "position": start_pos,
                    "length": len(para),
                    "is_heading": is_heading,
                    "heading_level": heading_level if is_heading else 0,
                    "contains_table": contains_table,
                    "contains_image": contains_image
                }
                paragraphs.append(paragraph)
            
            start_pos += len(para) + 2  # +2 是为了考虑分隔符 \n\n
        
        return paragraphs
    
    def file_to_markdown(self, file_path: str) -> str:
        """
        将不同格式的文件转换为Markdown格式
        
        参数:
            file_path (str): 文件路径
            
        返回:
            str: Markdown格式的文本内容
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)
            
            logger.info(f"开始将文件转换为Markdown: {file_path}, 类型: {file_ext}")
            
            # 根据文件类型调用不同的转换方法
            if file_ext == '.txt':
                return self._txt_to_markdown(file_path)
            elif file_ext == '.docx':
                return self._docx_to_markdown(file_path)
            elif file_ext == '.doc':
                return self._doc_to_markdown(file_path)
            elif file_ext == '.pdf':
                return self._pdf_to_markdown(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
                
        except Exception as e:
            error_message = f"转换文件到Markdown时出错: {str(e)}"
            logger.error(error_message)
            logger.error(traceback.format_exc())
            raise ValueError(error_message)
    
    def _txt_to_markdown(self, file_path: str) -> str:
        """
        将TXT文件转换为Markdown格式
        
        参数:
            file_path (str): TXT文件路径
            
        返回:
            str: Markdown格式的文本
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            # 对于TXT文件，我们只需要做简单的格式调整：
            # 1. 确保段落之间有空行
            # 2. 识别可能的标题并添加Markdown标记
            
            # 分割段落
            paragraphs = re.split(r'\n\s*\n', text)
            
            markdown_text = ""
            for i, para in enumerate(paragraphs):
                para = para.strip()
                if not para:
                    continue
                    
                # 检测可能的标题（短且独立的行）
                lines = para.split('\n')
                markdown_lines = []
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    # 判断是否可能是标题（长度较短的单行文本）
                    if len(lines) == 1 and len(line) < 100 and i > 0:
                        # 根据位置和前后文判断标题级别
                        if i == 0 or (i > 0 and len(paragraphs[i-1]) > 200):  # 可能是主标题
                            markdown_lines.append(f"# {line}")
                        else:  # 可能是子标题
                            markdown_lines.append(f"## {line}")
                    else:
                        markdown_lines.append(line)
                
                # 合并处理后的行
                if markdown_lines:
                    markdown_text += "\n\n" + "\n".join(markdown_lines)
            
            return markdown_text.strip()
            
        except Exception as e:
            logger.error(f"转换TXT到Markdown失败: {str(e)}")
            raise ValueError(f"转换TXT到Markdown失败: {str(e)}")
    
    def _docx_to_markdown(self, file_path: str) -> str:
        """
        将DOCX文件转换为Markdown格式
        
        参数:
            file_path (str): DOCX文件路径
            
        返回:
            str: Markdown格式的文本
        """
        try:
            document = Document(file_path)
            
            # 收集文档中的所有元素
            elements = []
            position_map = {}
            current_position = 0
            
            # 第一步：获取文档中所有段落和表格的逻辑位置，建立一个内部ID到位置的映射
            # 这对于准确定位表格非常重要
            all_items = []
            
            # 收集所有内容项（段落和表格）
            for item in document.element.body.iterchildren():
                if item.tag.endswith('p'):  # 段落
                    all_items.append(('paragraph', item))
                elif item.tag.endswith('tbl'):  # 表格
                    all_items.append(('table', item))
            
            # 为每个内容项分配一个位置标识符
            for idx, (item_type, item) in enumerate(all_items):
                position_map[item] = idx
            
            # 第二步：处理段落，使用位置映射
            for para_idx, para in enumerate(document.paragraphs):
                if not para.text.strip():
                    continue
                    
                # 获取段落在文档中的位置
                para_position = position_map.get(para._element, para_idx)
                
                # 获取段落样式
                style_name = para.style.name.lower() if para.style and para.style.name else ""
                
                # 处理标题
                if "heading" in style_name:
                    # 提取标题级别
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 1  # 默认为一级标题
                    
                    # 生成Markdown标题
                    heading_marks = "#" * min(level, 6)  # Markdown最高支持6级标题
                    elements.append({
                        "type": "heading",
                        "level": level,
                        "content": f"{heading_marks} {para.text.strip()}\n",
                        "position": para_position
                    })
                
                # 处理列表
                elif para.paragraph_format and para.paragraph_format.left_indent:
                    indent = para.paragraph_format.left_indent
                    # 估算列表级别
                    list_level = min(int(indent // 20), 3) if indent else 0
                    
                    if list_level > 0:
                        # 使用- 作为列表标记
                        elements.append({
                            "type": "list_item",
                            "level": list_level,
                            "content": para.text.strip(),
                            "position": para_position
                        })
                    else:
                        # 应用文本格式（粗体、斜体）
                        formatted_text = self._format_docx_text_runs(para)
                        elements.append({
                            "type": "paragraph",
                            "content": formatted_text,
                            "position": para_position
                        })
                
                # 处理普通段落
                else:
                    # 应用文本格式（粗体、斜体）
                    formatted_text = self._format_docx_text_runs(para)
                    elements.append({
                        "type": "paragraph",
                        "content": formatted_text,
                        "position": para_position
                    })
            
            # 第三步：处理表格，使用位置映射确保正确位置
            for table_idx, table in enumerate(document.tables):
                # 获取表格在文档中的真实位置
                table_position = position_map.get(table._element, current_position + table_idx)
                
                # 提取表格内容
                markdown_table = []
                
                # 处理表头（第一行）
                if table.rows:
                    header_row = []
                    for cell in table.rows[0].cells:
                        header_row.append(cell.text.strip() or " ")
                    markdown_table.append("| " + " | ".join(header_row) + " |")
                    
                    # 添加分隔行
                    markdown_table.append("| " + " | ".join(["---"] * len(header_row)) + " |")
                
                    # 处理数据行
                    for i, row in enumerate(table.rows):
                        if i == 0:  # 跳过表头行
                            continue
                            
                        row_cells = []
                        for cell in row.cells:
                            row_cells.append(cell.text.strip() or " ")
                        markdown_table.append("| " + " | ".join(row_cells) + " |")
                
                # 将表格添加到元素列表，标记表格前添加一个空行
                elements.append({
                    "type": "table",
                    "content": "\n" + "\n".join(markdown_table) + "\n\n",
                    "position": table_position
                })
            
            # 第四步：按位置排序所有元素
            elements.sort(key=lambda x: x["position"])
            
            # 第五步：合并元素，生成最终的Markdown内容
            markdown_text = ""
            current_list_level = 0
            in_list = False
            
            for element in elements:
                if element["type"] == "heading":
                    if in_list:
                        markdown_text += "\n"
                        in_list = False
                    markdown_text += "\n\n" + element["content"]
                
                elif element["type"] == "paragraph":
                    if in_list:
                        markdown_text += "\n"
                        in_list = False
                    markdown_text += element["content"] + "\n\n"
                
                elif element["type"] == "list_item":
                    level = element["level"]
                    
                    if not in_list or level != current_list_level:
                        markdown_text += "\n"
                    
                    # 使用- 作为列表标记
                    markdown_text += "  " * (level - 1) + "- " + element["content"] + "\n"
                    in_list = True
                    current_list_level = level
                
                elif element["type"] == "table":
                    if in_list:
                        markdown_text += "\n"
                        in_list = False
                    markdown_text += element["content"]
            
            return markdown_text.strip()
            
        except Exception as e:
            logger.error(f"转换DOCX到Markdown失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"转换DOCX到Markdown失败: {str(e)}")
    
    def _format_docx_text_runs(self, paragraph):
        """为DOCX段落中的文本添加Markdown格式化标记"""
        result = ""
        for run in paragraph.runs:
            text = run.text
            
            # 应用格式
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"__{text}__"
            
            result += text
        
        return result
    
    def _doc_to_markdown(self, file_path: str) -> str:
        """
        将DOC文件转换为Markdown格式
        
        参数:
            file_path (str): DOC文件路径
            
        返回:
            str: Markdown格式的文本
        """
        try:
            # 对于DOC文件，我们首先尝试转换为DOCX，然后再处理
            # 尝试主要方法：使用pywin32（仅在Windows环境下有效）
            try:
                # 首先检查系统是否为Windows
                if sys.platform.startswith('win'):
                    import win32com.client
                    
                    # 创建临时文件路径
                    temp_path = os.path.join(os.path.dirname(file_path), f"{uuid.uuid4()}.docx")
                    
                    # 使用Word应用程序转换doc为docx
                    word = None
                    doc = None
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        
                        # 尝试以只读方式打开文档
                        try:
                            doc = word.Documents.Open(file_path, ReadOnly=True)
                        except Exception:
                            doc = word.Documents.Open(file_path)
                        
                        # 保存为docx
                        doc.SaveAs(temp_path, 16)  # 16表示docx格式
                        doc.Close()
                        
                        # 检查文件是否实际创建
                        if os.path.exists(temp_path) and os.path.getsize(temp_path) > 100:
                            # 转换成功，使用DOCX转换方法
                            markdown_text = self._docx_to_markdown(temp_path)
                            
                            # 清理临时文件
                            try:
                                if os.path.exists(temp_path):
                                    os.remove(temp_path)
                            except Exception as cleanup_error:
                                logger.warning(f"清理临时文件失败: {str(cleanup_error)}")
                                
                            return markdown_text
                    finally:
                        # 确保文档被关闭
                        if doc:
                            try:
                                doc.Close(SaveChanges=False)
                            except:
                                pass
                        
                        # 确保Word应用被关闭
                        if word:
                            try:
                                word.Quit()
                            except:
                                pass
            except Exception as e:
                logger.warning(f"使用pywin32转换DOC失败: {str(e)}")
            
            # 备选方法：直接提取文本然后转换
            try:
                # 尝试使用docx2txt
                try:
                    import docx2txt
                    content = docx2txt.process(file_path)
                    
                    # 使用提取的文本生成简单的Markdown
                    paragraphs = content.split('\n\n')
                    markdown_text = ""
                    
                    for i, para in enumerate(paragraphs):
                        para = para.strip()
                        if not para:
                            continue
                            
                        # 检测可能的标题（短且独立的段落）
                        if len(para) < 100 and i > 0:
                            # 根据位置和前后文判断标题级别
                            if i == 0 or (i > 0 and len(paragraphs[i-1]) > 200):  # 可能是主标题
                                markdown_text += f"\n\n# {para}\n"
                            else:  # 可能是子标题
                                markdown_text += f"\n\n## {para}\n"
                        else:
                            markdown_text += f"\n\n{para}\n"
                    
                    return markdown_text.strip()
                    
                except ImportError:
                    logger.warning("docx2txt不可用，尝试其他方法")
                
                # 尝试使用textract
                try:
                    import textract
                    content = textract.process(file_path).decode('utf-8')
                    
                    # 使用提取的文本生成简单的Markdown
                    paragraphs = content.split('\n\n')
                    markdown_text = ""
                    
                    for i, para in enumerate(paragraphs):
                        para = para.strip()
                        if not para:
                            continue
                            
                        # 检测可能的标题（短且独立的段落）
                        if len(para) < 100 and i > 0:
                            # 根据位置和前后文判断标题级别
                            if i == 0 or (i > 0 and len(paragraphs[i-1]) > 200):  # 可能是主标题
                                markdown_text += f"\n\n# {para}\n"
                            else:  # 可能是子标题
                                markdown_text += f"\n\n## {para}\n"
                        else:
                            markdown_text += f"\n\n{para}\n"
                    
                    return markdown_text.strip()
                    
                except ImportError:
                    logger.warning("textract不可用，尝试其他方法")
                
                # 尝试使用antiword
                try:
                    import subprocess
                    result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    content = result.stdout.decode('utf-8')
                    
                    # 使用提取的文本生成简单的Markdown
                    paragraphs = content.split('\n\n')
                    markdown_text = ""
                    
                    for i, para in enumerate(paragraphs):
                        para = para.strip()
                        if not para:
                            continue
                            
                        # 检测可能的标题（短且独立的段落）
                        if len(para) < 100 and i > 0:
                            # 根据位置和前后文判断标题级别
                            if i == 0 or (i > 0 and len(paragraphs[i-1]) > 200):  # 可能是主标题
                                markdown_text += f"\n\n# {para}\n"
                            else:  # 可能是子标题
                                markdown_text += f"\n\n## {para}\n"
                        else:
                            markdown_text += f"\n\n{para}\n"
                    
                    return markdown_text.strip()
                    
                except (ImportError, FileNotFoundError):
                    logger.warning("antiword不可用，尝试其他方法")
                
                # 以文本方式直接读取，可能包含一些垃圾字符
                content = ""
                try:
                    with open(file_path, 'rb') as f:
                        data = f.read()
                    # 尝试从二进制数据中提取ASCII文本
                    for byte in data:
                        if 32 <= byte <= 126:  # 基本ASCII可打印字符
                            content += chr(byte)
                        elif byte in [9, 10, 13]:  # 制表符、换行符、回车符
                            content += chr(byte)
                    
                    # 清理内容，删除重复的空白字符
                    content = re.sub(r'\s+', ' ', content)
                    
                    # 尝试分割为段落
                    paragraphs = re.split(r'[\r\n]{2,}', content)
                    markdown_text = ""
                    
                    for i, para in enumerate(paragraphs):
                        para = para.strip()
                        if not para:
                            continue
                            
                        # 检测可能的标题（短且独立的段落）
                        if len(para) < 100 and i > 0:
                            # 根据位置和前后文判断标题级别
                            if i == 0 or (i > 0 and len(paragraphs[i-1]) > 200):  # 可能是主标题
                                markdown_text += f"\n\n# {para}\n"
                            else:  # 可能是子标题
                                markdown_text += f"\n\n## {para}\n"
                        else:
                            markdown_text += f"\n\n{para}\n"
                    
                    if len(markdown_text.strip()) > 100:  # 确保至少有一些有意义的文本
                        return markdown_text.strip()
                        
                except Exception as e:
                    logger.warning(f"二进制读取失败: {str(e)}")
                
                # 所有方法都失败，返回错误提示
                raise ValueError("无法处理此DOC文件，所有转换方法均已失败")
                
            except Exception as e:
                logger.error(f"转换DOC到Markdown失败: {str(e)}")
                raise ValueError(f"转换DOC到Markdown失败: {str(e)}")
                
        except Exception as e:
            logger.error(f"转换DOC到Markdown失败: {str(e)}")
            raise ValueError(f"转换DOC到Markdown失败: {str(e)}")
    
    def _pdf_to_markdown(self, file_path: str) -> str:
        """
        将PDF文件转换为Markdown格式
        
        参数:
            file_path (str): PDF文件路径
            
        返回:
            str: Markdown格式的文本
        """
        try:
            # 使用PyMuPDF打开PDF
            doc = fitz.open(file_path)
            markdown_text = ""
            
            # 添加文档标题
            file_name = os.path.basename(file_path)
            markdown_text += f"# {os.path.splitext(file_name)[0]}\n\n"
            
            # 添加文档元数据摘要
            try:
                metadata = doc.metadata
                if metadata:
                    markdown_text += "## 文档信息\n\n"
                    if metadata.get("title"):
                        markdown_text += f"- **标题**: {metadata.get('title')}\n"
                    if metadata.get("author"):
                        markdown_text += f"- **作者**: {metadata.get('author')}\n"
                    if metadata.get("subject"):
                        markdown_text += f"- **主题**: {metadata.get('subject')}\n"
                    if metadata.get("creator"):
                        markdown_text += f"- **创建者**: {metadata.get('creator')}\n"
                    if metadata.get("producer"):
                        markdown_text += f"- **生成器**: {metadata.get('producer')}\n"
                    markdown_text += "\n"
            except Exception as e:
                logger.warning(f"提取PDF元数据时出错: {str(e)}")
            
            # 添加目录（如果存在）
            try:
                toc = doc.get_toc()
                if toc:
                    markdown_text += "## 目录\n\n"
                    for level, title, page in toc:
                        indent = "  " * (level - 1)
                        markdown_text += f"{indent}- [{title}](#page-{page})\n"
                    markdown_text += "\n"
            except Exception as e:
                logger.warning(f"提取PDF目录时出错: {str(e)}")
            
            # 逐页处理PDF内容
            for page_idx, page in enumerate(doc):
                # 添加页码标记
                page_num = page_idx + 1
                markdown_text += f"## 第{page_num}页 {{#page-{page_num}}}\n\n"
                
                try:
                    # 收集页面上的所有内容，包括文本块、表格和图片
                    page_elements = []
                    
                    # 提取页面上的文本块（带位置信息）
                    page_dict = page.get_text("dict")
                    blocks = page_dict.get("blocks", [])
                    
                    # 处理页面上的文本块
                    for block_idx, block in enumerate(blocks):
                        block_type = block.get("type", 0)
                        
                        # 处理文本块
                        if block_type == 0:  # 文本块
                            lines = block.get("lines", [])
                            if not lines:
                                continue
                            
                            # 提取文本块的位置信息
                            bbox = block.get("bbox", [0, 0, 0, 0])
                            y_pos = bbox[1]  # 块的顶部Y坐标
                            
                            # 合并文本块中的所有行
                            text_content = ""
                            for line in lines:
                                spans = line.get("spans", [])
                                for span in spans:
                                    text_content += span.get("text", "")
                                text_content += " "
                            
                            text_content = text_content.strip()
                            if not text_content:
                                continue
                            
                            # 判断文本块的类型（标题、列表、普通段落）
                            if len(text_content) < 100 and text_content.strip().endswith((":", "：")) and not text_content.strip().startswith(("•", "-", "*")):
                                # 可能是标题（短文本且以冒号结尾）
                                page_elements.append({
                                    "type": "heading",
                                    "content": text_content,
                                    "position": y_pos,
                                    "level": 3  # 默认为三级标题
                                })
                            elif text_content.strip().startswith(("•", "■", "◆", "▪", "○", "●", "-", "*", "1.", "2.", "3.")):
                                # 可能是列表项
                                # 清理列表标记
                                text_content = re.sub(r'^[•■◆▪○●\-*]\s*|^\d+\.\s*', '- ', text_content)
                                page_elements.append({
                                    "type": "list_item",
                                    "content": text_content,
                                    "position": y_pos
                                })
                            else:
                                # 普通段落
                                page_elements.append({
                                    "type": "paragraph",
                                    "content": text_content,
                                    "position": y_pos
                                })
                        
                        # 处理图片块 (PDF中的图片块类型通常为1)
                        elif block_type == 1:
                            # 获取图片块的位置
                            bbox = block.get("bbox", [0, 0, 0, 0])
                            y_pos = bbox[1]  # 图片的顶部Y坐标
                            
                            # 添加图片标记
                            img_content = f"\n![图 {page_idx+1}-{block_idx}](图片在原PDF中的第{page_idx+1}页位置Y={y_pos:.1f})\n\n"
                            
                            page_elements.append({
                                "type": "image",
                                "content": img_content,
                                "position": y_pos
                            })
                    
                    # 尝试提取表格
                    try:
                        tables = page.find_tables()
                        if tables and hasattr(tables, 'tables') and tables.tables:
                            for table_idx, table in enumerate(tables.tables):
                                try:
                                    # 获取表格的位置
                                    y_pos = 0
                                    if hasattr(table, 'rect') and table.rect:
                                        y_pos = table.rect.y0
                                    
                                    # 提取表格数据
                                    rows = []
                                    cols = 0
                                    
                                    if hasattr(table, 'cells') and table.cells:
                                        # 确定表格的行和列
                                        cell_list = list(table.cells)
                                        if cell_list:
                                            rows_dict = {}
                                            for cell in cell_list:
                                                row_idx = cell.y0
                                                if row_idx not in rows_dict:
                                                    rows_dict[row_idx] = []
                                                rows_dict[row_idx].append((cell.x0, cell.text.strip() or " "))
                                            
                                            # 排序行并构建表格数据
                                            sorted_rows = [rows_dict[k] for k in sorted(rows_dict.keys())]
                                            rows = []
                                            for row in sorted_rows:
                                                sorted_cells = [cell[1] for cell in sorted(row, key=lambda x: x[0])]
                                                rows.append(sorted_cells)
                                                cols = max(cols, len(sorted_cells))
                                    
                                    # 创建Markdown表格
                                    if rows and cols > 0:
                                        # 确保所有行具有相同的列数
                                        for i in range(len(rows)):
                                            while len(rows[i]) < cols:
                                                rows[i].append(" ")
                                        
                                        # 构建表格内容
                                        table_content = f"\n**表 {page_idx+1}-{table_idx+1}:**\n\n"
                                        table_content += "| " + " | ".join(rows[0]) + " |\n"
                                        table_content += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
                                        
                                        # 添加数据行
                                        for row in rows[1:]:
                                            table_content += "| " + " | ".join(row) + " |\n"
                                        
                                        table_content += "\n"
                                        
                                        # 将表格添加到页面元素
                                        page_elements.append({
                                            "type": "table",
                                            "content": table_content,
                                            "position": y_pos
                                        })
                                except Exception as e:
                                    logger.warning(f"处理PDF表格时出错 (页 {page_idx+1}, 表 {table_idx+1}): {str(e)}")
                    except Exception as e:
                        logger.warning(f"查找PDF页面表格时出错 (页 {page_idx+1}): {str(e)}")
                    
                    # 按Y坐标排序所有元素（从上到下）
                    page_elements.sort(key=lambda x: x["position"])
                    
                    # 合并处理后的元素成Markdown
                    page_markdown = ""
                    current_section = None
                    prev_element_type = None
                    
                    for element in page_elements:
                        element_type = element["type"]
                        
                        # 处理不同类型的元素
                        if element_type == "heading":
                            # 更新当前章节
                            current_section = element["content"]
                            page_markdown += f"### {element['content']}\n\n"
                        elif element_type == "paragraph":
                            # 如果前一个元素不是段落，确保有足够的空行
                            if prev_element_type != "paragraph":
                                page_markdown += "\n"
                            page_markdown += f"{element['content']}\n\n"
                        elif element_type == "list_item":
                            page_markdown += f"{element['content']}\n"
                        elif element_type == "table":
                            # 表格内容已包含换行，所以这里不需要额外添加
                            page_markdown += element["content"]
                        elif element_type == "image":
                            page_markdown += element["content"]
                        
                        prev_element_type = element_type
                    
                    # 添加生成的页面内容到总文档
                    markdown_text += page_markdown
                    
                except Exception as e:
                    logger.error(f"处理PDF页面时出错 (页 {page_idx+1}): {str(e)}")
                    logger.error(traceback.format_exc())
                    markdown_text += f"*无法处理此页面内容: {str(e)}*\n\n"
            
            return markdown_text.strip()
            
        except Exception as e:
            logger.error(f"转换PDF到Markdown失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"转换PDF到Markdown失败: {str(e)}")

    def _read_file(self, file_path: str) -> str:
        """读取文件内容"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()


# if __name__ == "__main__":
#     # 测试代码
#     file_handler = FileToMarkdown()
#     try:
#         # 测试转换TXT文件
#         # result = file_handler.file_to_markdown("test.txt")
#         # print(result)
        
#         # 测试转换DOCX文件
#         result = file_handler.file_to_markdown("D:\桌面\培训资料\Docs\技术方案\密码应用方案\密码监管平台建设方案.docx")
#         print(result)
        
#         # 测试转换PDF文件
#         # result = file_handler.file_to_markdown("test.pdf")
#         # print(result)
        
#         print("文件转换测试完成")
#     except Exception as e:
#         print(f"测试失败: {str(e)}") 